import re
import unicodedata
from typing import Any
from src.documentos_en_word.programa import Programa
from docx.document import Document
from docx.shared import Inches


class Contenido:

    def __init__(self, data: Any):
        if "rich_text" in data and data["rich_text"]:
            self._contenido_plano = "".join(
                fragment.get("text", {}).get("content", "")
                for fragment in data["rich_text"]
            )
        else:
            self._contenido_plano = ""

    def insertar_en_documento(self, documento: Document):
        """Por defecto, inserta como párrafo simple."""
        documento.add_paragraph(self._contenido_plano)


class Encabezado(Contenido):
    def __init__(self, data: Any):
        super().__init__(data)
        self._contenido_plano = (
            self._normalizar_titulo_si_pertenece_a_una_seccion_clave()
        )

    def _normalizar_titulo_si_pertenece_a_una_seccion_clave(self):
        titulo = self._contenido_plano
        # Diccionario de secciones con sus palabras clave principales
        secciones_clave_y_variantes = {
            "1- OBJETIVOS/ EXPECTATIVAS DE LOGROS/ CAPACIDADES": [
                r"objetivos?",
                r"expectativas?",
                r"logros?",
                r"capacidades?",
            ],
            "2- CONTENIDOS": [r"contenidos?"],
            "3- METODOLOGÍA DE TRABAJO": [r"metodolog[ií]a", r"trabajo"],
            "4- CRITERIO DE EVALUACIÓN": [r"criterio[s]?", r"evaluaci[oó]n"],
            "5- BIBLIOGRAFÍA PARA EL ESTUDIANTE": [
                r"bibliograf[ií]a",
                r"estudiante[s]?",
            ],
        }

        # Normalizamos el texto de entrada
        titulo_normalizado = titulo.strip().lower()

        for seccion_clave, variantes in secciones_clave_y_variantes.items():
            # Coincidencia exacta ignorando mayúsculas/minúsculas y tildes
            if self._eliminar_acentos(titulo_normalizado) in [
                self._eliminar_acentos(parte_de_una_seccion_clave.lower())
                for parte_de_una_seccion_clave in re.split(
                    r"/|\|", seccion_clave.split("-", 1)[1]
                )
            ]:
                return seccion_clave
            # Coincidencia por palabra clave
            for variante in variantes:
                if re.search(variante, titulo_normalizado, re.IGNORECASE):
                    return seccion_clave

        return titulo

    def _eliminar_acentos(self, s: str) -> str:
        return "".join(
            c
            for c in unicodedata.normalize("NFD", s)
            if unicodedata.category(c) != "Mn"
        )


class Encabezado1(Encabezado):
    def insertar_en_documento(self, documento: Document):
        documento.add_heading(self._contenido_plano.upper(), level=1)


class Encabezado2(Encabezado):
    def insertar_en_documento(self, documento: Document):
        documento.add_heading(self._contenido_plano.upper(), level=2)


class Encabezado3(Encabezado):
    def insertar_en_documento(self, documento: Document):
        documento.add_heading(self._contenido_plano, level=3)


class Parrafo(Contenido):
    def insertar_en_documento(self, documento: Document):
        documento.add_paragraph(self._contenido_plano)


class ListaConViñetas(Contenido):
    def insertar_en_documento(self, documento: Document):
        # python-docx no soporta listas automáticas, pero se puede simular con símbolos
        documento.add_paragraph(f"• {self._contenido_plano}", style="List Bullet")


class ListaNumerada(Contenido):
    def insertar_en_documento(self, documento: Document):
        documento.add_paragraph(self._contenido_plano, style="List Number")


class ListaConToDo(Contenido):

    def __init__(self, data: Any):
        super().__init__(data)
        self.checked = data.get("checked", False)

    def insertar_en_documento(self, documento: Document):
        check = "☑" if self.checked else "☐"
        documento.add_paragraph(f"{check} {self._contenido_plano}")


class Llamada(Contenido):  # callout
    def insertar_en_documento(self, documento: Document):
        # Simular un callout usando un símbolo
        documento.add_paragraph(f"💡 {self._contenido_plano}")


class Cite(Contenido):  # quote
    def insertar_en_documento(self, documento: Document):
        p = documento.add_paragraph(self._contenido_plano)
        # Simular cita usando sangría
        p.paragraph_format.left_indent = Inches(0.5)
        p.style = "Intense Quote" if "Intense Quote" in documento.styles else p.style


class Codigo(Contenido):

    def __init__(self, data: Any):
        super().__init__(data)
        self._lenguaje = data.get("language", "")

    def insertar_en_documento(self, documento: Document):
        # Simular bloque de código con fuente monoespaciada
        p = documento.add_paragraph()
        run = p.add_run(self._contenido_plano)
        run.font.name = "Courier New"
        # Opcional: agregar sombreado o estilo personalizado si se requiere


class Separador(Contenido):  # divider
    def __init__(self, _: Any):
        self._contenido_plano = "---"

    def insertar_en_documento(self, documento: Document):
        # Simular divisor con una línea
        p = documento.add_paragraph()
        run = p.add_run("_" * 30)
        run.italic = True


class BloqueDeContenido:
    _fabrica_de_contenido = dict(
        {
            "paragraph": lambda bloque: Parrafo(bloque),
            "heading_1": lambda bloque: Encabezado1(bloque),
            "heading_2": lambda bloque: Encabezado2(bloque),
            "heading_3": lambda bloque: Encabezado3(bloque),
            "bulleted_list_item": lambda bloque: ListaConViñetas(bloque),
            "numbered_list_item": lambda bloque: ListaNumerada(bloque),
            "to_do": lambda bloque: ListaConToDo(bloque),
            "callout": lambda bloque: Llamada(bloque),
            "quote": lambda bloque: Cite(bloque),
            "code": lambda bloque: Codigo(bloque),
            "divider": lambda bloque: Separador(bloque),
        }
    )

    def __init__(self, data: Any):
        tipo = data["type"]
        bloque = data[tipo]

        if tipo in BloqueDeContenido._fabrica_de_contenido:
            self._contenido = BloqueDeContenido._fabrica_de_contenido[tipo](bloque)
        else:
            self._contenido = Contenido(bloque)

    def insertar_en_documento(self, programa: Programa):
        self._contenido.insertar_en_documento(programa._documento)
        return self
