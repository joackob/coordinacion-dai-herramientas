from typing import Any
from src.documentos_en_word.programa import Programa
from docx.document import Document
from docx.shared import Inches, Pt


class Contenido:
    _contenido_plano: str

    def __init__(self, data: Any):
        if "rich_text" in data and data["rich_text"]:
            self._contenido_plano = "".join(
                fragment.get("text", {}).get("content", "")
                for fragment in data["rich_text"]
            )
        else:
            self._contenido_plano = ""

    def insertar_en_documento(self, documento: Document):
        """Por defecto, inserta como párrafo simple."""
        documento.add_paragraph(self._contenido_plano)


class Encabezado1(Contenido):
    def insertar_en_documento(self, documento: Document):
        titulo = documento.add_heading(self._contenido_plano, level=1)
        if titulo.style:
            titulo.style.font.name = "Arial"
            titulo.style.font.size = Pt(14)
            titulo.style.font.bold = True


class Encabezado2(Contenido):
    def insertar_en_documento(self, documento: Document):
        documento.add_heading(self._contenido_plano, level=2)


class Encabezado3(Contenido):
    def insertar_en_documento(self, documento: Document):
        documento.add_heading(self._contenido_plano, level=3)


class Parrafo(Contenido):
    def insertar_en_documento(self, documento: Document):
        documento.add_paragraph(self._contenido_plano)


class ListaConViñetas(Contenido):
    def insertar_en_documento(self, documento: Document):
        # python-docx no soporta listas automáticas, pero se puede simular con símbolos
        documento.add_paragraph(f"• {self._contenido_plano}", style="List Bullet")


class ListaNumerada(Contenido):
    def insertar_en_documento(self, documento: Document):
        documento.add_paragraph(self._contenido_plano, style="List Number")


class ListaConToDo(Contenido):
    checked: bool

    def __init__(self, data: Any):
        super().__init__(data)
        self.checked = data.get("checked", False)

    def insertar_en_documento(self, documento: Document):
        check = "☑" if self.checked else "☐"
        documento.add_paragraph(f"{check} {self._contenido_plano}")


class Llamada(Contenido):  # callout
    def insertar_en_documento(self, documento: Document):
        # Simular un callout usando un símbolo
        documento.add_paragraph(f"💡 {self._contenido_plano}")


class Cite(Contenido):  # quote
    def insertar_en_documento(self, documento: Document):
        p = documento.add_paragraph(self._contenido_plano)
        # Simular cita usando sangría
        p.paragraph_format.left_indent = Inches(0.5)
        p.style = "Intense Quote" if "Intense Quote" in documento.styles else p.style


class Codigo(Contenido):
    lenguaje: str

    def __init__(self, data: Any):
        super().__init__(data)
        self.lenguaje = data.get("language", "")

    def insertar_en_documento(self, documento: Document):
        # Simular bloque de código con fuente monoespaciada
        p = documento.add_paragraph()
        run = p.add_run(self._contenido_plano)
        run.font.name = "Courier New"
        # Opcional: agregar sombreado o estilo personalizado si se requiere


class Separador(Contenido):  # divider
    def __init__(self, _: Any):
        self._contenido_plano = "---"

    def insertar_en_documento(self, documento: Document):
        # Simular divisor con una línea
        p = documento.add_paragraph()
        run = p.add_run("_" * 30)
        run.italic = True


class BloqueDeContenido:
    _contenido: Contenido

    def __init__(self, data: Any):
        tipo = data["type"]
        bloque = data[tipo]

        if tipo == "paragraph":
            self._contenido = Parrafo(bloque)
        elif tipo == "heading_1":
            self._contenido = Encabezado1(bloque)
        elif tipo == "heading_2":
            self._contenido = Encabezado2(bloque)
        elif tipo == "heading_3":
            self._contenido = Encabezado3(bloque)
        elif tipo == "bulleted_list_item":
            self._contenido = ListaConViñetas(bloque)
        elif tipo == "numbered_list_item":
            self._contenido = ListaNumerada(bloque)
        elif tipo == "to_do":
            self._contenido = ListaConToDo(bloque)
        elif tipo == "callout":
            self._contenido = Llamada(bloque)
        elif tipo == "quote":
            self._contenido = Cite(bloque)
        elif tipo == "code":
            self._contenido = Codigo(bloque)
        elif tipo == "divider":
            self._contenido = Separador(bloque)
        else:
            self._contenido = Contenido(bloque)

    def insertar_en_documento(self, programa: Programa):
        self._contenido.insertar_en_documento(programa._documento)
        return self
