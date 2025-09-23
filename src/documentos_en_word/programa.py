from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.styles.style import ParagraphStyle
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from config import ubicacion_de_documento_plantilla
from config import ubicacion_carpeta_donde_guardar_programas_generados


class Programa:
    _jefe_de_departamento: str = "Andrés Navarro"
    _anio_ciclo: dict[str, str] = dict(
        {
            "3ro": "3ro - Ciclo Intermedio",
            "4to": "4to - Ciclo Intermedio",
            "5to": "5to - Ciclo Superior",
            "6to": "6to - Ciclo Superior",
        }
    )

    def __init__(self, asignatura: str, anio: str, carga_horaria: int, area: str):
        self._docentes = list[str]()
        # Abrir el documento de Word que servirá como plantilla
        ubicacion_de_documento_plantilla.resolve()
        self._documento = Document(f"{ubicacion_de_documento_plantilla}")

        # Modificar estilos del encamezado de nivel 1
        estilo_encabezado_1 = self._documento.styles["Heading 1"]
        if isinstance(estilo_encabezado_1, ParagraphStyle):
            estilo_encabezado_1.font.name = "Arial"
            estilo_encabezado_1.font.size = Pt(14)
            estilo_encabezado_1.font.bold = True
            estilo_encabezado_1.paragraph_format.space_before = Pt(14)
            estilo_encabezado_1.paragraph_format.space_after = Pt(8)

        estilo_parrafo = self._documento.styles["Normal"]
        if isinstance(estilo_parrafo, ParagraphStyle):
            estilo_parrafo.font.name = "Arial"
            estilo_parrafo.font.size = Pt(12)
            estilo_parrafo.paragraph_format.line_spacing = 1.5

        # Agregar estilos que no existen en el documento
        estilo_encabezado_2 = self._documento.styles.add_style(
            "Heading 2", WD_STYLE_TYPE.PARAGRAPH
        )
        if isinstance(estilo_encabezado_2, ParagraphStyle):
            estilo_encabezado_2.font.name = "Arial"
            estilo_encabezado_2.font.size = Pt(14)
            estilo_encabezado_2.font.bold = True
            estilo_encabezado_2.paragraph_format.space_before = Pt(14)
            estilo_encabezado_2.paragraph_format.space_after = Pt(8)

        estilo_encabezado_3 = self._documento.styles.add_style(
            "Heading 3", WD_STYLE_TYPE.PARAGRAPH
        )
        if isinstance(estilo_encabezado_3, ParagraphStyle):
            estilo_encabezado_3.font.name = "Arial"
            estilo_encabezado_3.font.size = Pt(13)
            estilo_encabezado_3.font.bold = True
            estilo_encabezado_3.paragraph_format.space_before = Pt(14)
            estilo_encabezado_3.paragraph_format.space_after = Pt(8)

        estilo_item_de_lista_desordenada = self._documento.styles.add_style(
            "List Bullet", WD_STYLE_TYPE.PARAGRAPH
        )
        if isinstance(estilo_item_de_lista_desordenada, ParagraphStyle):
            estilo_item_de_lista_desordenada.font.name = "Arial"
            estilo_item_de_lista_desordenada.font.size = Pt(12)
            estilo_item_de_lista_desordenada.paragraph_format.line_spacing = 1.5
            estilo_item_de_lista_desordenada.paragraph_format.left_indent = Pt(36)

        # Obtener la tabla que contiene los datos requeridos
        self._tabla_con_datos_requeridos = self._documento.tables[0]

        # Completar los datos básicos en la tabla del documento
        # Asignatura
        asignatura_en_tabla = self._tabla_con_datos_requeridos.cell(1, 1)
        asignatura_en_tabla.text = asignatura
        # Año/Ciclo
        anio_ciclo_en_tabla = self._tabla_con_datos_requeridos.cell(2, 1)
        anio_ciclo_en_tabla.text = Programa._anio_ciclo[anio]
        # Campo de formación
        campo_de_formacion_en_tabla = self._tabla_con_datos_requeridos.cell(3, 1)
        campo_de_formacion_en_tabla.text = f"Técnico Especifico - Área de {area}"
        # Carga horaria
        carga_horaria_en_tabla = self._tabla_con_datos_requeridos.cell(4, 1)
        carga_horaria_en_tabla.text = f"{carga_horaria} horas cátedra"
        # Jefe de departamento
        jefe_de_departamento_en_tabla = self._tabla_con_datos_requeridos.cell(5, 1)
        jefe_de_departamento_en_tabla.text = Programa._jefe_de_departamento
        # Dar estilos requeridos a la celdas que fueron modificadas
        for celda in [
            asignatura_en_tabla,
            anio_ciclo_en_tabla,
            campo_de_formacion_en_tabla,
            carga_horaria_en_tabla,
            jefe_de_departamento_en_tabla,
        ]:
            for parrafo in celda.paragraphs:
                parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in parrafo.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(14)
                    run.font.bold = True

    def agregar_nombres_de_docentes(self, nombres: list[str]):
        self._docentes.extend(nombres)
        docentes_en_tabla = self._tabla_con_datos_requeridos.cell(6, 1)
        docentes_en_tabla = docentes_en_tabla.paragraphs[0].add_run(
            ", ".join(self._docentes)
        )
        docentes_en_tabla.font.name = "Arial"
        docentes_en_tabla.font.size = Pt(14)
        docentes_en_tabla.font.bold = True
        return self

    def separar_tabla_con_datos_del_contenido(self):
        # Insertar un salto de página despues de la primera pagina
        self._documento.add_page_break()

    def guardar(self):
        ubicacion_final_documento = (
            ubicacion_carpeta_donde_guardar_programas_generados
            / f"{self._tabla_con_datos_requeridos.cell(1, 1).text.lower().replace(' ', '_')}.docx"
        )
        ubicacion_final_documento.resolve()
        self._documento.save(f"{ubicacion_final_documento}")
        return self
