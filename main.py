from notion_client import Client
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# CONFIGURA ESTOS VALORES:
NOTION_API_KEY = "tu-integracion-secreta"
DATABASE_ID = "tu-database-id"
LOGO_PATH = "logo.png"  # Debe ser PNG, JPG o BMP

# Inicializa Notion
notion = Client(auth=NOTION_API_KEY)


# Función para obtener el contenido de la página como texto plano
def extract_text_from_blocks(blocks):
    texts = []
    for block in blocks["results"]:
        t = block.get(block["type"], {}).get("text", [])
        if t:
            texts.append("".join([part["plain_text"] for part in t]))
    return "\n".join(texts)


# Obtén todas las páginas de la base de datos
def get_database_pages(database_id):
    results = []
    next_cursor = None
    while True:
        query = {"database_id": database_id}
        if next_cursor:
            query["start_cursor"] = next_cursor
        resp = notion.databases.query(**query)
        results.extend(resp["results"])
        if resp.get("has_more"):
            next_cursor = resp["next_cursor"]
        else:
            break
    return results


# Obtiene los bloques de contenido de una página
def get_page_blocks(page_id):
    blocks = notion.blocks.children.list(page_id)
    return blocks


# Inserta imagen en encabezado y pie de página
def insert_logo_in_header_footer(section, logo_path):
    for header in [section.header, section.footer]:
        p = header.paragraphs[0]
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(1))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# Crea el documento Word para una materia
def create_word_doc(materia, content_text, logo_path, docentes):
    doc = Document()
    section = doc.sections[0]
    insert_logo_in_header_footer(section, logo_path)

    # Tabla de información
    table = doc.add_table(rows=7, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Programa"
    table.cell(0, 1).text = ""
    table.cell(1, 0).text = "Asignatura"
    table.cell(1, 1).text = materia["nombre"]
    table.cell(2, 0).text = "Año - Ciclo"
    table.cell(2, 1).text = materia["año"]
    table.cell(3, 0).text = "Campo de Formación"
    table.cell(3, 1).text = "Técnico Especifico - Especialidad TICs"
    table.cell(4, 0).text = "Carga horaria semanal"
    table.cell(4, 1).text = str(materia["carga_horaria"])
    table.cell(5, 0).text = "Jefe/a de departamento"
    table.cell(5, 1).text = "Andrés Navarro"
    table.cell(6, 0).text = "Docente"
    table.cell(6, 1).text = ", ".join(docentes)

    # Salto de página
    doc.add_page_break()
    # Contenido de la materia
    doc.add_paragraph(content_text)
    return doc


# MAIN
def main():
    pages = get_database_pages(DATABASE_ID)
    print(f"Se encontraron {len(pages)} materias.")
    for page in pages:
        props = page["properties"]
        nombre = props["Nombre"]["title"][0]["plain_text"]
        año = props["Año"]["select"]["name"]
        carga_horaria = props["Carga horaria semanal"]["number"]
        # DOCENTES debe ser una propiedad multi-select o relation de la página
        docentes = []
        if "Docentes" in props and props["Docentes"]["people"]:
            docentes = [d["name"] for d in props["Docentes"]["people"]]
        else:
            docentes = ["(Sin asignar)"]

        materia = {"nombre": nombre, "año": año, "carga_horaria": carga_horaria}
        page_id = page["id"]
        blocks = get_page_blocks(page_id)
        content_text = extract_text_from_blocks(blocks)

        doc = create_word_doc(materia, content_text, LOGO_PATH, docentes)
        # Guarda el archivo
        filename = f"{nombre}.docx"
        doc.save(filename)
        print(f"Guardado: {filename}")


if __name__ == "__main__":
    main()
